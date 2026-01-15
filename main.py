# Imports

import os
import tempfile
import streamlit as st
from io import BytesIO
import hashlib
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import (
    CSVLoader,
    EverNoteLoader,
    PyMuPDFLoader,
    TextLoader,
    UnstructuredEmailLoader,
    UnstructuredHTMLLoader,
    UnstructuredMarkdownLoader,
    UnstructuredEPubLoader,
    UnstructuredODTLoader,
    UnstructuredPowerPointLoader,
    UnstructuredWordDocumentLoader,
    UnstructuredExcelLoader,
    UnstructuredImageLoader,
    SRTLoader
)

from streamlit_option_menu import option_menu
from langchain_groq import ChatGroq
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import StrOutputParser
from tavily import TavilyClient
import requests
from langchain_core.output_parsers import JsonOutputParser
from PIL import Image
# Environment setup

SERPAPI_API_KEY=st.secrets["SERPAPI_API_KEY"]
SERPER_API_KEY=st.secrets["SERPER_API_KEY"]
os.environ["GROQ_API_KEY"] = st.secrets["GROQ_API_KEY"]
llm = ChatGroq(model="llama-3.1-8b-instant")
TAVILY_API_KEY=st.secrets["TAVILY_API_KEY"]
class Config:
    SERPAPI_API_KEY = SERPAPI_API_KEY
    SERPER_API_KEY = SERPER_API_KEY

serperai_api_key = Config.SERPER_API_KEY

# ------------------------
# CONFIG
# ------------------------
Loader_map = {
    ".csv": (CSVLoader, {}),
    ".enex": (EverNoteLoader, {}),
    ".pdf": (PyMuPDFLoader, {}),
    ".txt": (TextLoader, {"encoding": "utf8"}),
    ".eml": (UnstructuredEmailLoader, {}),
    ".html": (UnstructuredHTMLLoader, {}),
    ".md": (UnstructuredMarkdownLoader, {}),
    ".mmd": (UnstructuredMarkdownLoader, {}),
    ".epub": (UnstructuredEPubLoader, {}),
    ".odt": (UnstructuredODTLoader, {}),
    ".ppt": (UnstructuredPowerPointLoader, {}),
    ".pptx": (UnstructuredPowerPointLoader, {}),
    ".doc": (UnstructuredWordDocumentLoader, {}),
    ".docx": (UnstructuredWordDocumentLoader, {}),
    ".xlsx": (UnstructuredExcelLoader, {}),
    ".xls": (UnstructuredExcelLoader, {}),
    ".jpg": (UnstructuredImageLoader, {}),
    ".png": (UnstructuredImageLoader, {}),
    ".srt": (SRTLoader, {}),
}

text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=3000,
    chunk_overlap=300,
    separators=["\n\n", "\n", " ", ""]
)


# ------------------------
# FILE GENERATORS
# ------------------------

def find_youtube_videos_tool(missing_skill: str):
    """Searches for YouTube tutorials for a specific missing skill."""
    if not missing_skill:
        return []

    skill = missing_skill.replace(" ", "+")
    url = f"https://google.serper.dev/videos?q={skill}&apiKey={serperai_api_key}"

    response = requests.get(url)
    videos = response.json().get("videos", [])

    clean_results = []
    for video in videos[:5]:
        clean_results.append({
            "title": video.get("title"),
            "video": video.get("link"),
            "thumbnail": video.get("imageUrl", "")
        })

    return clean_results

def generate_pdf(text: str) -> BytesIO:
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer)
    styles = getSampleStyleSheet()
    story = [Paragraph(line, styles["Normal"]) for line in text.split("\n")]
    doc.build(story)
    buffer.seek(0)
    return buffer

def generate_docx(text: str) -> BytesIO:
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def get_loader(file_path: str):
    st.session_state.setdefault("final_summary", None)
    st.session_state.setdefault("executive_summary", None)

    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    if ext in Loader_map:
        cls, args = Loader_map[ext]
        return cls(file_path, **args)
    raise ValueError(f"Unsupported file extension: {ext}")

# ------------------------
# LLM
# ------------------------

def load_resume_with_langchain(uploaded_file):
    """
    Load resume text from almost any file type using LangChain loaders.
    """

    suffix = os.path.splitext(uploaded_file.name)[1].lower()

    # Save file temporarily (required by LangChain loaders)
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
        tmp_file.write(uploaded_file.getbuffer())
        tmp_path = tmp_file.name

    try:
        loader= Loader_map.get(suffix)
        loader_class, loader_args = loader if loader else (TextLoader, {"encoding": "utf8"})
        loader = loader_class(tmp_path, **loader_args)
        documents = loader.load()

        # Merge all pages into one string
        resume_text = "\n".join(doc.page_content for doc in documents)

        return resume_text.strip()
    except Exception as e:
        return f"Error loading resume: {e}"

    finally:
        os.remove(tmp_path)

# Helper functions and tools
def tavily_search_tool(query):
    client = TavilyClient(TAVILY_API_KEY)
    response = client.search(
        query=query,
        include_answer="advanced",
        search_depth="advanced"
    )
    tav_answer=response["answer"]

    for i in response["results"]:
        tav_answer += f"\n{i['content']}\n"
    return response

def save_report_tool(filename: str, content: str):
    """
    Saves the final research report to a Word document.
    Args:
        filename: The name of the file (e.g., 'bitcoin_analysis.docx')
        content: The full text content to write into the document.
    """
    try:
        doc = Document()
        doc.add_heading('Research Report', 0)
        
        # Split content by newlines and parse markdown
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                continue
                
            # Handle Headers
            if line.startswith('# '):
                doc.add_heading(line[2:].strip(), level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:].strip(), level=3)
            # Handle Bullet points
            elif line.startswith('- ') or line.startswith('* '):
                # Check for bold text in bullets
                clean_line = line[2:].strip()
                p = doc.add_paragraph(style='List Bullet')
                if '**' in clean_line:
                    parts = clean_line.split('**')
                    for i, part in enumerate(parts):
                        run = p.add_run(part)
                        if i % 2 == 1: # Odd parts were between ** **
                            run.bold = True
                else:
                    p.add_run(clean_line)
            # Handle Standard Paragraphs
            else:
                p = doc.add_paragraph()
                if '**' in line:
                    parts = line.split('**')
                    for i, part in enumerate(parts):
                        run = p.add_run(part)
                        if i % 2 == 1:
                            run.bold = True
                else:
                    p.add_run(line)
                
        doc.save(filename)
        return f"File saved successfully: {filename}"
    except Exception as e:
        return f"Error saving file: {str(e)}"
    


# ------------------------
# PROMPTS
# ------------------------

career_prompt = PromptTemplate(
    input_variables=["resume", "job_description"],
    template="""
You are a career coach AI. You are given a **resume** and a **job description**. Perform the following tasks:

1. Compare the resume to the job description and calculate a **match score** (0-100%).
2. Identify the **top missing skill** if the match score is less than 100%.
3. Identify **potential contacts** mentioned in the job description.
4. Generate **cold email templates** for these contacts.

Return the output in this exact JSON format:
{{
  "name": "<Candidate Name from Resume>",
  "reason": "<Explain why the score is what it is>",
  "match_score": 85,
  "top_missing_skill": "Skill Name",
  "potential_contacts": ["HR Manager", "Team Lead"],
  "cold_emails": [
    {{"HR": "Dear HR, I am reaching out because..."}},
    {{"Team Lead": "Dear Team Lead, I am reaching out because..."}}
  ]
}}

**Rules:**
- All fields must be present
- Do not include any extra text outside JSON

Resume:
{resume}

Job Description:
{job_description}
"""
)

career_chain = career_prompt | llm | JsonOutputParser()


writer_prompt = PromptTemplate(
    template="""
You are a **Technical Writer**.

Your task is to summarize the research text **strictly in alignment with the given topic**.
All sections must stay focused on this topic. Do not introduce unrelated concepts.

---

## Topic
{topic}

---

## Instructions
Summarize the research text below into a **professional, detailed, and comprehensive report**.

Structure the report using the following **Markdown headers exactly**:

# Executive Summary
# Key Findings
# Detailed Analysis
# Implications
# Future Outlook

### Writing Guidelines
- Keep the narrative tightly centered on the topic
- Use bullet points where appropriate
- Use **bold text** for emphasis
- Write clear, professional paragraphs
- Ensure the content is **rich, insightful, and not scanty**
- If the research text is broad, **filter and reframe** it to fit the topic

---

## Research Text
{text}
""",
    input_variables=["topic", "text"]
)

writer_chain = writer_prompt | llm | StrOutputParser()

chunk_summary_prompt = PromptTemplate(
    template="""
You are an expert document analyst.
Summarize the following content in detail.
Use bullet points where helpful.
{text}
""",
    input_variables=["text"]
)

final_summary_prompt = PromptTemplate(
    template="""
You are a professional technical writer.
Combine the summaries into a coherent document summary.
Use headings and subheadings.
Start with a short overview.
{text}
""",
    input_variables=["text"]
)

executive_summary_prompt = PromptTemplate(
    template="""
Create a concise executive summary (max 5 bullet points).
Focus on key insights and implications.
{text}
""",
    input_variables=["text"]
)

chunk_summary_chain = chunk_summary_prompt | llm | StrOutputParser()
final_summary_chain = final_summary_prompt | llm | StrOutputParser()
executive_summary_chain = executive_summary_prompt | llm | StrOutputParser()

# ------------------------
# STREAMLIT UI
# ------------------------
from streamlit_option_menu import option_menu

selected = option_menu(
    menu_title=None,
    options=[
        "Home",
        "Doc Summarizer",
        "Researcher Agent",
        "Resume Agent"
    ],
    icons=[
        "house-fill",          # Home
        "file-earmark-text",   # Document summarizer
        "search",              # Research agent
        "person-lines-fill"    # Resume agent
    ],
    default_index=0,
    orientation="horizontal"
)

if selected == "Home":
    col1, col2 = st.columns(2)

    with col1:
        st.title("🧠 AI-Powered Knowledge & Career Assistant")
        st.caption(
            """
            <b>
            1. Upload documents and get instant, structured summaries — executive or detailed.<br><br>
            2. Run deep AI-powered research on any topic and receive professional reports with insights and outlooks.<br><br>
            3. Match your resume to job descriptions, uncover missing skills, and generate smart cold emails.<br><br>
            4. Export summaries, research, and career analysis as PDF, DOCX, or Markdown — ready to share.
            </b>
            """,
            unsafe_allow_html=True
        )

    with col2:
        image = Image.open("./Frame_4.jpg")
        # image = image.resize((600, 750))
        st.image(image)


elif selected == "Doc Summarizer":
    st.title("📄 Document Summarizer")

    # Session state
    st.session_state.setdefault("final_summary", None)
    st.session_state.setdefault("executive_summary", None)
    st.session_state.setdefault("uploaded_filename", None)

    summary_mode = st.radio(
        "Summary Type",
        ["Detailed Summary", "Executive Summary"],
        horizontal=True
    )

    if st.button("🔄 Clear Summary"):
        st.session_state.final_summary = None
        st.session_state.executive_summary = None
        st.rerun()

    uploaded_file = st.file_uploader(
        "Upload a document",
        type=[ext.replace(".", "") for ext in Loader_map.keys()]
    )
    # ✅ Detect new file upload
    if uploaded_file:
        if st.session_state.uploaded_filename != uploaded_file.name:
            st.session_state.uploaded_filename = uploaded_file.name
            st.session_state.final_summary = None
            st.session_state.executive_summary = None
        file_bytes = uploaded_file.getvalue()
        file_hash = hashlib.md5(file_bytes).hexdigest()

        if st.session_state.get("uploaded_hash") != file_hash:
            st.session_state.uploaded_hash = file_hash
            st.session_state.final_summary = None
            st.session_state.executive_summary = None
        
    if uploaded_file:
        suffix = os.path.splitext(uploaded_file.name)[1]

        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(uploaded_file.read())
            tmp_path = tmp.name

        try:
            loader = get_loader(tmp_path)
            docs = loader.load()
            full_text = "\n".join(d.page_content for d in docs).strip()

            if len(full_text) < 500:
                st.warning("Document too short to summarize.")
            else:
                chunks = text_splitter.split_text(full_text)
                st.info(f"Split into {len(chunks)} sections")

                if st.button("Generate Summary"):
                    progress = st.progress(0)
                    section_summaries = []

                    for i, chunk in enumerate(chunks):
                        section_summaries.append(
                            chunk_summary_chain.invoke({"text": chunk})
                        )
                        progress.progress((i + 1) / len(chunks))

                    combined_text = "\n\n".join(section_summaries)

                    if summary_mode == "Executive Summary":
                        st.session_state.executive_summary = executive_summary_chain.invoke(
                            {"text": combined_text}
                        )
                        st.session_state.final_summary = None
                    else:
                        st.session_state.final_summary = final_summary_chain.invoke(
                            {"text": combined_text}
                        )
                        st.session_state.executive_summary = None

        except Exception as e:
            st.error(f"Error: {e}")

        finally:
            os.remove(tmp_path)

    # ------------------------
    # OUTPUT
    # ------------------------

    summary_text = (
        st.session_state.executive_summary
        or st.session_state.final_summary
    )

    if st.session_state.executive_summary:
        st.subheader("🧠 Executive Summary")
        st.markdown(st.session_state.executive_summary)

    if st.session_state.final_summary:
        st.subheader("📘 Detailed Document Summary")
        st.markdown(st.session_state.final_summary)

    if summary_text:
        st.divider()
        col1, col2, col3 = st.columns(3)

        with col1:
            st.download_button(
                "⬇️ Markdown",
                f"# Summary\n\n{summary_text}",
                "summary.md",
                "text/markdown"
            )

        with col2:
            st.download_button(
                "⬇️ PDF",
                generate_pdf(summary_text),
                "summary.pdf",
                "application/pdf"
            )

        with col3:
            st.download_button(
                "⬇️ DOCX",
                generate_docx(summary_text),
                "summary.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        st.success("Summary generated successfully!")
elif selected == "Researcher Agent":
    st.title("🚀 Groq Research Agent")
    st.caption("Writer Agent Final Report")
    topic = st.text_input(
    "What should I research?",
    placeholder="e.g., Nvidia Stock Performance 2024"
    )

    start = st.button("Start Research")



    if start and topic:
        st.info("Running research workflow... ⚡")
        answer=tavily_search_tool(topic)
        st.success("Research completed! 🎉")
        agent_answer=writer_chain.invoke({"topic":topic,"text":answer})
        st.success("Report generated! 📄")
        st.markdown(agent_answer)
        save_report_tool(f"{topic}.docx", agent_answer)
        st.success(f"Report saved as {topic}.docx! 💾")
elif selected == "Resume Agent":
    st.title("📄 AI Resume–Job Matching & Career Coach")
    st.caption("Analyze resume fit, missing skills & generate cold emails")
    resume_upload=st.radio(
        "Upload Resume Method",
        ["File Upload", "Paste Text"],
        horizontal=True
    )
    if resume_upload=="File Upload":
         uploaded_resume = st.file_uploader(
            "Upload your resume",
            type=[ext.replace(".", "") for ext in Loader_map.keys()],
            accept_multiple_files=False
        )
         resume_text=""
         if uploaded_resume:
             resume_text = load_resume_with_langchain(uploaded_resume)
             if resume_text.startswith("Error loading resume:"):
                 st.error(resume_text)
             else:
                 st.success("Resume loaded successfully!")
                 st.text_area("Resume Text", resume_text, height=200)
    else:
        resume_text = st.text_area(
            "Paste your resume text here",
            placeholder="e.g., Experienced Software Engineer with expertise in Python, Java, and cloud technologies..."
        )

    job_description_text = st.text_area(
        "Paste the job description here",
        placeholder="e.g., Software Engineer at TechCorp"
    )

    analyze_btn = st.button("🔍 Analyze Career Match", type="primary")

    # ---------------- ANALYSIS ----------------

    if analyze_btn:
        if not resume_text or not job_description_text:
            st.warning("Please provide both resume and job description.")
        else:
            with st.spinner("Analyzing your profile..."):
                try:
                    result = career_chain.invoke({
                        "resume": resume_text,
                        "job_description": job_description_text
                    })

                    # Fetch learning resources
                    result["learning_plan"] = find_youtube_videos_tool(
                        result.get("top_missing_skill")
                    )

                    # ---------------- DISPLAY RESULTS ----------------

                    st.success("Analysis Complete ✅")

                    st.subheader("👤 Candidate")
                    st.write(result["name"])

                    st.subheader("📊 Match Score")
                    st.progress(result["match_score"] / 100)
                    st.write(f"**{result['match_score']}% match**")

                    st.subheader("🧠 Reasoning")
                    st.write(result["reason"])

                    st.subheader("❌ Top Missing Skill")
                    st.warning(result["top_missing_skill"])

                    st.subheader("📞 Potential Contacts")
                    st.write(", ".join(result["potential_contacts"]))

                    st.subheader("✉️ Cold Email Templates")
                    for email in result["cold_emails"]:
                        for role, content in email.items():
                            with st.expander(f"{role} Email"):
                                st.write(content)

                    st.subheader("📚 Learning Resources")
                    if result["learning_plan"]:
                        for video in result["learning_plan"]:
                            st.markdown(f"**{video['title']}**")
                            st.markdown(f"[Watch Video]({video['video']})")
                    else:
                        st.info("No learning resources found.")

                    with st.expander("🧾 Raw JSON Output", expanded=False):
                        st.json(result)

                except Exception as e:
                    st.error(f"Error: {e}")